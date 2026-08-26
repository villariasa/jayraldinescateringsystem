import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas

# Define paths
MD_PATH = "Catering_Present/WORKSHEET3-SYSTEM-REQUEST.md"
DOCX_PATH = "Catering_Present/WORKSHEET3-SYSTEM-REQUEST.docx"
PDF_PATH = "Catering_Present/WORKSHEET3-SYSTEM-REQUEST.pdf"

# Content Data
TITLE = "PC 317 — Systems Analysis and Design"
SUBTITLE = "Worksheet No. 3: System Request"
SYSTEM_TITLE = "SYSTEM REQUEST: JAYRALDINE'S CATERING MANAGEMENT SYSTEM"

HEADER_INFO = {
    "Course / Subject": "PC 317 — Systems Analysis and Design",
    "Worksheet No.": "Worksheet No. 3 (System Request)",
    "Project Title": "Jayraldine's Catering Management System (Desktop Hub & Mobile Tablet Client)",
    "Proponents / Student Names": "BSIT Capstone Project Group (Jayraldine's Catering System Team)",
    "Target Organization": "Jayraldine's Catering Services",
    "Date": "Academic Year 2025-2026"
}

SECTIONS = [
    {
        "heading": "1. Project Sponsor",
        "content": [
            "• Primary Sponsor: Business Owner / Management of Jayraldine's Catering Services.",
            "• Technical Proponents: BSIT Capstone Student Project Proponents.",
            "• Lead Contact Person: Business Owner in coordination with the Lead Student System Analyst / Developer."
        ]
    },
    {
        "heading": "2. Business Need",
        "content": [
            "This project was initiated to replace Jayraldine's Catering's inefficient, paper-based manual operations and fragmented tools (Excel spreadsheets, paper logs, phone messaging apps) with a centralized, automated Catering Management & Mobile Tablet System.",
            "Key operational challenges prompting this system request include:",
            "• Overbooking & Capacity Risks: Lack of automated safeguards led to scheduling conflicts and overbooking beyond daily kitchen/pax limits (max daily pax capacity).",
            "• Revenue Leakage & Uncollected Payments: Absence of strict downpayment enforcement before booking confirmation exposed the business to cancelled events and unpaid receivables.",
            "• Kitchen Coordination Bottlenecks: Kitchen staff relied on vague text notes without structured, per-dish itemized task checklists, causing preparation errors and delivery delays.",
            "• Manual Client Communications: Confirmations and receipts were generated manually, leading to missed client updates, delayed billing, and lack of professional record-keeping.",
            "• Financial Invisibility: Business tracking focused strictly on gross revenue without logging operational expenses (food cost, labor, transport, utilities), preventing accurate net profit calculation.",
            "• Slow Off-site Client Intake: Inability to capture bookings and display package options interactively during client consultations away from the main office."
        ]
    },
    {
        "heading": "3. Business Requirements",
        "content": [
            "The system delivers a comprehensive dual-client ecosystem (Desktop PC Hub + Mobile Tablet Client) with the following core business capabilities currently implemented:",
            "A. Dual-Client Architecture & Data Synchronization:",
            "• Desktop PC Hub (PySide6 + PostgreSQL/SQLite): Full administrative control, deep financial analytics, comprehensive audit logs, DB backup/restore, and master data management.",
            "• Mobile/Tablet POS Client (PySide6 / Android APK): Touch-first order entry wizard designed for fast customer registration, interactive package/menu selection, dynamic charge/discount calculations, and digital terms acceptance.",
            "• Offline Capability & Offline-to-PC Sync: Tablet operates standalone using local SQLite storage. Supports master data import (.db or .xlsx) from PC and seamless 1-click database merging into the main PC hub without overwriting or duplicating existing records.",
            "B. Booking & Capacity Enforcement:",
            "• Complete Booking Lifecycle: Manage bookings from Pending → Confirmed → In Progress → Completed → Cancelled.",
            "• Automated Capacity Hard Block: Automatically checks total booked guest count against daily limits (default max 600 pax) and blocks new bookings if capacity is exceeded.",
            "• Downpayment Enforcement: Enforces a configurable minimum downpayment percentage (default 30%) before a booking status can be set to Confirmed.",
            "• Cancellation Audit: Captures structured cancellation reasons for reporting and quality control.",
            "C. Interactive Kitchen Kanban & Dish Task Checklists:",
            "• Kitchen Kanban Board: Visual order tracking across states (Queued → Preparing → In Progress → Ready → Delivered → Done).",
            "• Per-Dish Task Checklist: Itemized task breakdown (kitchen_tasks) auto-generated per dish, allowing kitchen staff to check off preparation steps in real time.",
            "D. Billing, Invoices & Financial Management:",
            "• Automated Invoicing: Real-time balance calculations (Unpaid, Partial, Paid) with instant receipt and invoice generation.",
            "• Multi-Category Expense Logging: Records operational costs under categorized heads (Food Cost, Labor, Transport, Utilities, Equipment, Other).",
            "• Net Profit Analytics: Dashboard KPI metrics displaying real-time Net Profit (Revenue minus Expenses).",
            "E. Multi-Channel Client Notifications:",
            "• Branded Email Notifications (SMTP): Automated transmission of PDF invoices and formal payment receipts directly to client email addresses.",
            "• SMS Notifications (Semaphore API): Instant automated SMS booking confirmations sent to client mobile numbers.",
            "• In-App Event Alerts: Built-in notification bell panel and toast popups alerting staff 24 hours, 30 minutes, and at event start time.",
            "F. Customer Relationship Management (CRM) & Loyalty Program:",
            "• Customer Loyalty Tiers: Automatic customer classification into Bronze (1–2 events), Silver (3–5), Gold (6–9), and VIP (10+ events) tiers based on booking history.",
            "• Customer Follow-Up Reminders: Task scheduler for client outreach and anniversary/re-booking reminders.",
            "G. Terms & Conditions Compliance:",
            "• Digital Terms Acknowledgment: Integrated contract terms wizard with mandatory customer digital signature/acknowledgment and schema versioning (ta_version).",
            "H. System Administration, Security & Auditing:",
            "• Comprehensive Audit Trail: Logs user actions, affected tables, timestamps, and JSON snapshots of pre/post modifications (audit_logs).",
            "• Single-Click Database Backup & Restore: Built-in utility to export and restore database SQL dumps (pg_dump/psql) for disaster recovery."
        ]
    },
    {
        "heading": "4. Business Value",
        "content": [
            "The implemented system provides high tangible and intangible value to Jayraldine's Catering:",
            "Tangible Value:",
            "• 100% Elimination of Double-Bookings & Overbooking: Hard pax capacity constraints prevent booking beyond operational capacity.",
            "• Guaranteed Upfront Cash Flow: 30% downpayment enforcement eliminates unpaid client cancellations and no-show financial risks.",
            "• 60% Faster Client Booking Intake: Touch-optimized Tablet wizard cuts order placement time from 20+ minutes down to 5-8 minutes.",
            "• Complete Profit Visibility: Real-time net profit tracking isolates high-cost operational categories and improves gross margins.",
            "• Paper & Printing Cost Reduction: Digital PDF invoice generation, email dispatch, and SMS confirmations significantly reduce physical paperwork costs.",
            "Intangible Value:",
            "• Enhanced Professional Brand Image: Automated instant SMS confirmations, branded PDF receipts, and tablet-based order taking boost customer trust.",
            "• Improved Kitchen Coordination & Accuracy: Itemized dish checklists reduce kitchen preparation errors and missing menu items during events.",
            "• Stronger Customer Retention: Automated loyalty tier badges and follow-up reminders encourage repeat client bookings.",
            "• Total Accountability & Data Security: Comprehensive audit logs and single-click DB backup guarantee data integrity and operational oversight."
        ]
    },
    {
        "heading": "5. Special Issues or Constraints",
        "content": [
            "• Project Context & Scope: Developed as an academic BSIT Capstone Project tailored specifically to the operational workflow of Jayraldine's Catering Services.",
            "• Technical Environment: Hybrid architecture utilizing Python 3.11, PySide6 (Qt for Python), PostgreSQL for the main PC hub, and standalone SQLite for the mobile/tablet client.",
            "• Third-Party Dependencies: Automated notifications rely on active SMTP internet connection for email dispatch and Semaphore API credentials for SMS dispatch.",
            "• Hardware Compatibility: Desktop PC management hub runs on standard Windows/Linux PCs; Mobile POS app builds to standalone Android APKs and touch-screen tablets.",
            "• Academic Review: Requires formal evaluation and approval by the Capstone Defense Panel per university academic standards."
        ]
    }
]

def generate_markdown():
    print("Generating Markdown file...")
    md_content = []
    md_content.append(f"# {TITLE}")
    md_content.append(f"## {SUBTITLE}\n")
    
    for k, v in HEADER_INFO.items():
        md_content.append(f"**{k}:** {v}")
    
    md_content.append("\n---\n")
    md_content.append(f"## {SYSTEM_TITLE}\n")
    
    for sec in SECTIONS:
        md_content.append(f"### {sec['heading']}\n")
        for line in sec['content']:
            md_content.append(line)
        md_content.append("")
        
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(md_content))
    print(f"Saved: {MD_PATH}")

def generate_docx():
    print("Generating Word DOCX file...")
    doc = docx.Document()
    
    # Page setup - standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Title Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(TITLE)
    run_t.font.size = Pt(18)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(SUBTITLE)
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Header Info Table
    table = doc.add_table(rows=len(HEADER_INFO), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Style table
    for i, (k, v) in enumerate(HEADER_INFO.items()):
        row = table.rows[i]
        
        # Key cell
        cell_k = row.cells[0]
        cell_k.width = Inches(2.2)
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k + ":")
        rk.font.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        # Value cell
        cell_v = row.cells[1]
        cell_v.width = Inches(4.5)
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        
        # Shading
        shading_xml = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "F2F5F9" if i % 2 == 0 else "FFFFFF"))
        cell_k._tc.get_or_add_tcPr().append(shading_xml)
        shading_xml2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "F2F5F9" if i % 2 == 0 else "FFFFFF"))
        cell_v._tc.get_or_add_tcPr().append(shading_xml2)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Main Document Header
    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sys = p_sys.add_run(SYSTEM_TITLE)
    run_sys.font.size = Pt(13)
    run_sys.font.bold = True
    run_sys.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    # Sections
    for sec in SECTIONS:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(14)
        p_h.paragraph_format.space_after = Pt(4)
        p_h.paragraph_format.keep_with_next = True
        rh = p_h.add_run(sec['heading'])
        rh.font.size = Pt(12)
        rh.font.bold = True
        rh.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        for line in sec['content']:
            p_c = doc.add_paragraph()
            p_c.paragraph_format.space_after = Pt(3)
            p_c.paragraph_format.line_spacing = 1.15
            
            if line.startswith("•"):
                p_c.paragraph_format.left_indent = Inches(0.25)
                r_bullet = p_c.add_run("• ")
                r_bullet.font.bold = True
                r_bullet.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
                rc = p_c.add_run(line[2:])
            elif line.startswith("A.") or line.startswith("B.") or line.startswith("C.") or line.startswith("D.") or line.startswith("E.") or line.startswith("F.") or line.startswith("G.") or line.startswith("H.") or line.startswith("Tangible") or line.startswith("Intangible"):
                p_c.paragraph_format.space_before = Pt(6)
                rc = p_c.add_run(line)
                rc.font.bold = True
                rc.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
            else:
                rc = p_c.add_run(line)
                
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#666666"))
        
        # Footer
        footer_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 36, footer_text)
        self.drawString(54, 36, "PC 317 Worksheet No. 3 — System Request | Jayraldine's Catering")
        
        # Decorative line above footer
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(54, 48, 612 - 54, 48)
        
        self.restoreState()

def generate_pdf():
    print("Generating PDF file...")
    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_title = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1B365D"),
        alignment=1, # Center
        spaceAfter=2
    )
    
    style_subtitle = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#4B6B94"),
        alignment=1, # Center
        spaceAfter=12
    )
    
    style_table_key = ParagraphStyle(
        'TableKey',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1B365D")
    )
    
    style_table_val = ParagraphStyle(
        'TableVal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1E293B")
    )
    
    style_sys_title = ParagraphStyle(
        'SysTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#1B365D"),
        spaceBefore=10,
        spaceAfter=8
    )
    
    style_heading = ParagraphStyle(
        'SecHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1B365D"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    style_subheading = ParagraphStyle(
        'SubHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#2C5282"),
        spaceBefore=6,
        spaceAfter=2,
        keepWithNext=True
    )
    
    style_body = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor("#334155"),
        spaceAfter=3
    )
    
    style_bullet = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor("#334155"),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )
    
    story = []
    
    # Title
    story.append(Paragraph(TITLE, style_title))
    story.append(Paragraph(SUBTITLE, style_subtitle))
    
    # Info Table
    table_data = []
    for k, v in HEADER_INFO.items():
        table_data.append([
            Paragraph(k + ":", style_table_key),
            Paragraph(v, style_table_val)
        ])
        
    t = Table(table_data, colWidths=[150, 354])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor("#F1F5F9"), colors.HexColor("#FFFFFF")]),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor("#1E293B")),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1B365D"), spaceAfter=10))
    
    # System Title
    story.append(Paragraph(SYSTEM_TITLE, style_sys_title))
    
    # Sections
    for sec in SECTIONS:
        story.append(Paragraph(sec['heading'], style_heading))
        for line in sec['content']:
            if line.startswith("•"):
                story.append(Paragraph(f"<font color='#1B365D'><b>•</b></font> {line[2:]}", style_bullet))
            elif line.startswith("A.") or line.startswith("B.") or line.startswith("C.") or line.startswith("D.") or line.startswith("E.") or line.startswith("F.") or line.startswith("G.") or line.startswith("H.") or line.startswith("Tangible") or line.startswith("Intangible"):
                story.append(Paragraph(line, style_subheading))
            else:
                story.append(Paragraph(line, style_body))
                
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Saved: {PDF_PATH}")

if __name__ == "__main__":
    generate_markdown()
    generate_docx()
    generate_pdf()
