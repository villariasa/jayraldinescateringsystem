import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

BASE_DIR = "/home/villarias/Projects/jayraldinescateringsystem"
MD_OUT = os.path.join(BASE_DIR, "PROJECT_PROPOSAL_AND_PRESENTATION_GUIDE.md")
DOCX_OUT = os.path.join(BASE_DIR, "PROJECT_PROPOSAL_AND_PRESENTATION_GUIDE.docx")
PDF_OUT = os.path.join(BASE_DIR, "PROJECT_PROPOSAL_AND_PRESENTATION_GUIDE.pdf")

MD_CONTENT = """# 🍽️ JAYRALDINE'S CATERING MANAGEMENT & SELF-SERVICE KIOSK SYSTEM
## Formal Project Proposal & Functional Specification Document

---

## 1. Project Background & Business Case

* **Client**: Jayraldine's Catering Services (Cebu, Philippines).
* **Project Classification**: Integrated Catering ERP & Self-Service Kiosk System.
* **System Scope**: Dual-Platform Solution (Windows PC Back-Office Management Suite + 100% Offline Tablet/Mobile Field Kiosk).
* **Project Goal**: Automate and modernize end-to-end catering operations — replacing manual paperwork with automated client ordering, live financial estimations, downpayment enforcement, kitchen preparation workflows, expense tracking, AI business forecasting, and comprehensive reporting.

---

## 2. Operational Problems vs. Proposed System Solutions

| Current Manual Operational Challenge | Proposed System Solution |
| :--- | :--- |
| **Manual Headcount & Math Calculation Errors** | Automated mathematical engine computing exact base totals, per-pax rates, custom add-ons, downpayment percentages, and balance ledgers without human error. |
| **Unmonitored Expenses & Blind Profit Margins** | Integrated Cost of Goods Sold (COGS), labor, transport, and utility expense tracking deducting directly from gross revenue to calculate true Net Profit. |
| **Double Bookings & Venue Capacity Overload** | Automated event date conflict detection and daily capacity hard-blocking (e.g. max 600 pax/day). |
| **Remote Venue Disconnectivity** | 100% Offline-First architecture running self-contained SQLite engines on local hardware with zero internet dependency. |
| **Consultation & Ordering Bottlenecks** | Self-service guided touch ordering flow reducing client consultation time from 45 minutes to under 5 minutes. |

---

## 3. Proposal Topic Breakdown for 3 Presenters

---

### 📌 PRESENTER 1: Project Scope, System Architecture & Executive Oversight

#### Core Topics to Present:
* **Dual-Platform System Architecture**:
  * **PC Desktop Workstation (`Catering_Present`)**: Executive control center for administration, master catalog setup, invoicing, kitchen operations, and financial auditing.
  * **Portable Field Kiosk (`Tablet_PWA` / APK)**: Standalone, offline client ordering terminal used during bridal fairs, food tastings, and on-site consultations.
  * **2-Way Synchronization**: Seamless data sync between field tablets and back-office computers via master Excel templates and SQLite database exports.
* **Executive Dashboard & Real-Time Business KPIs**:
  * Live monitoring of Total Gross Revenue, Total Operating Expenses, Net Profit Margins, and Active Bookings.
  * Pipeline tracking categorized by Inquiries, Confirmed Events, In-Preparation, Completed, and Cancelled bookings.
* **Capacity Governance & Scheduling Protection**:
  * Automated capacity validation preventing overbooking beyond kitchen production capacity limits.
* **Customer CRM & Loyalty Engine**:
  * Client booking history tracking with automated loyalty tiers (*Bronze, Silver, Gold, VIP*) and scheduled follow-up reminder logs.

---

### 📌 PRESENTER 2: Client Ordering Subsystem, Sales Workflows & Billing Automation

#### Core Topics to Present:
* **Guided 6-Step Client Ordering Engine**:
  * **Step 1 (Client Identification)**: Returning customer search + separated Cebu Barangay/City dropdown and street address fields.
  * **Step 2 (Event Schedule & Buffet Package)**: Touch occasion selector (*Weddings, Birthdays, Debuts, Corporate, Baptisms*) + **Package Details & Inclusions Viewer** displaying food rates and standard equipment inclusions (*chafing dishes, table skirting, servers, floral centerpieces*).
  * **Step 3 (Categorized Menu Selection)**: Structured food catalog grouped by dish types (*Beef, Pork, Chicken, Seafood, Pasta, Desserts, Beverages*) with sticky top category navigation.
  * **Step 4 (Add-ons & Equipment Upselling)**: One-tap add-on selections for high-margin items (*Whole Lechon platters, Dessert bars, Sound & Lighting rentals*).
  * **Step 5 (Billing & Legal Terms)**: Real-time deposit calculator, payment channel tracking (*Cash, GCash, Bank Transfer*), and legally binding terms acknowledgement.
  * **Step 6 (Pre-Booking Verification & Instant PDF Receipt)**: Verification summary before submission + automated offline itemized PDF receipt generation saved directly to device storage.
* **Billing, Invoicing & Notification System**:
  * Automated official commercial invoice generation (`INV-XXXX`), payment milestone tracking, and partial balance ledgers.
  * Automated booking confirmations and PDF invoices dispatched via Email (SMTP) and SMS (Semaphore API).
  * Minimum downpayment enforcement (e.g. 30%) blocking event confirmation without verified deposit.

---

### 📌 PRESENTER 3: Financial Management, Expense Tracking, AI Forecasting & Kitchen Operations

#### Core Topics to Present:
* **Cash Flow & Expense Tracking (Net Profit Calculation)**:
  * Real-time Cash Inflow vs. Outflow tracking with net liquidity balances.
  * Itemized expense recording categorized by Cost of Goods Sold (COGS), Food Ingredients, Labor/Waitstaff Wages, Fuel/Transport, Utility Overheads, and Equipment Rentals.
  * Automated formula: `Gross Revenue - Total Operating Expenses = True Net Profit`.
* **AI Intelligence & Business Forecasting**:
  * **AI Demand & Volume Forecasting**: Machine learning predictions on seasonal booking spikes, peak holiday months, and expected guest volume.
  * **Smart Ingredient & Portion Estimator**: Automated calculations predicting raw food material requirements based on pax count and selected dishes to prevent food waste.
  * **Menu Engineering Analytics**: Intelligent analysis highlighting top-selling high-margin items and underperforming dishes.
* **Kitchen Operations & Task Breakdown (Kanban Workflow)**:
  * Real-time preparation pipeline (*Pending, Cooking, Plating, Ready for Transport, Dispatched*).
  * Interactive per-dish task checklist ensuring standard recipe execution and zero missed menu items.
* **Audit Governance, Reporting & Data Privacy**:
  * Comprehensive analytics: Monthly/Annual revenue comparisons, package sales distribution, and payment method charts.
  * Complete audit logs recording actor identity, action type, timestamps, and JSON data states.
  * 100% offline data ownership with zero third-party cloud vulnerabilities and zero recurring monthly fees.

---

## 4. Complete System Functional Module Matrix

| Module | Core Functionality | Platform |
| :--- | :--- | :--- |
| **Executive Dashboard** | Real-time KPIs, Revenue, Net Profit, Capacity alerts, Recent Activity | PC Back-Office |
| **Booking & Kiosk Wizard** | 6-step guided ordering, package selection, downpayment calculation | PC + Tablet Kiosk |
| **Calendar & Scheduling** | Visual event timeline, date conflict detection, daily capacity load | PC Back-Office |
| **Menu & Package Catalog** | CRUD packages, pax price matrix, dish categories, photo management | PC + Tablet Kiosk |
| **Billing & Invoicing** | Invoice generation, downpayment enforcement, balance collection | PC + Tablet Kiosk |
| **Receipt & Notifications** | PDF receipt generation, SMTP email dispatcher, SMS gateway | PC + Tablet (PDF) |
| **Cash Flow Tracker** | Inflow/Outflow tracking, payment method breakdown, net balance | PC Back-Office |
| **Expense & COGS Manager** | Food costs, labor wages, logistics, utilities, Net Profit margin | PC Back-Office |
| **AI Demand & Forecaster** | Pax forecasting, ingredient quantity estimator, menu insights | PC Back-Office |
| **Kitchen Kanban & Checklist**| Visual preparation cards, per-dish task checklist, order dispatch | PC Back-Office |
| **Customer CRM & Loyalty** | Booking history, automated loyalty tiers (Bronze to VIP), follow-ups | PC + Tablet Kiosk |
| **Reports & Financial Analytics**| Revenue charts, package sales breakdown, PDF/Excel export | PC Back-Office |
| **Security, Audit & Backups** | Full change audit logs, 1-click DB backup/restore, Excel data sync | PC + Tablet Kiosk |

---

## 5. Technical Specifications & Architecture

* **Back-Office PC System**: Native Python/Qt Desktop Application compiled to Windows executable (`.exe`).
* **Tablet Kiosk App**: HTML5 / JavaScript / CSS + Client-Side SQLite WASM + Standalone Android APK (`.apk`, 7.1 MB).
* **Offline Capability**: 100% self-contained database engine; zero internet connection required for ordering, querying, or receipt printing.
* **Image Engine**: Full HD $1920 \times 1440$ bicubic image processor with WebP/JPEG compression.
* **Data Security & Privacy**: Stored locally on physical business hardware; zero third-party cloud data leaks or monthly subscription fees.
"""

def generate_all():
    # 1. Write Markdown file
    with open(MD_OUT, "w", encoding="utf-8") as f:
        f.write(MD_CONTENT)
    print(f"Generated Markdown: {MD_OUT}")

    # 2. Generate DOCX
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    title = doc.add_heading("JAYRALDINE'S CATERING MANAGEMENT & KIOSK SYSTEM", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(225, 29, 72)

    sub = doc.add_paragraph("Formal Project Proposal & Functional Specification Document")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Background
    h1 = doc.add_heading("1. Project Background & Business Case", level=1)
    for run in h1.runs: run.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.add_run("• Client: ").bold = True
    p.add_run("Jayraldine's Catering Services (Cebu, Philippines)\n")
    p.add_run("• Project Classification: ").bold = True
    p.add_run("Integrated Catering ERP & Self-Service Kiosk System\n")
    p.add_run("• System Scope: ").bold = True
    p.add_run("Dual-Platform Solution (Windows PC Back-Office Suite + 100% Offline Tablet/Mobile Kiosk)\n")
    p.add_run("• Project Goal: ").bold = True
    p.add_run("Automate and modernize end-to-end catering operations — replacing manual paperwork with automated client ordering, live financial estimations, downpayment enforcement, kitchen preparation workflows, expense tracking, AI business forecasting, and comprehensive reporting.")

    # 2. Problems vs Solutions Table
    doc.add_heading("2. Operational Problems vs. Proposed System Solutions", level=1)
    table_prob = doc.add_table(rows=1, cols=2)
    table_prob.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_p = table_prob.rows[0].cells
    hdr_p[0].text = "Current Manual Operational Challenge"
    hdr_p[1].text = "Proposed System Solution"
    for cell in hdr_p:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(r'<w:shd {} w:fill="E11D48"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)

    prob_data = [
        ("Manual Headcount & Math Calculation Errors", "Automated mathematical engine computing exact base totals, per-pax rates, custom add-ons, downpayment percentages, and balance ledgers without human error."),
        ("Unmonitored Expenses & Blind Profit Margins", "Integrated Cost of Goods Sold (COGS), labor, transport, and utility expense tracking deducting directly from gross revenue to calculate true Net Profit."),
        ("Double Bookings & Venue Capacity Overload", "Automated event date conflict detection and daily capacity hard-blocking (e.g. max 600 pax/day)."),
        ("Remote Venue Disconnectivity", "100% Offline-First architecture running self-contained SQLite engines on local hardware with zero internet dependency."),
        ("Consultation & Ordering Bottlenecks", "Self-service guided touch ordering flow reducing client consultation time from 45 minutes to under 5 minutes.")
    ]
    for ch, sol in prob_data:
        row_cells = table_prob.add_row().cells
        row_cells[0].text = ch
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = sol

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. Presenter Breakdown
    doc.add_heading("3. Proposal Topic Breakdown for 3 Presenters", level=1)

    # Presenter 1
    doc.add_heading("PRESENTER 1: Project Scope, System Architecture & Executive Oversight", level=2)
    p1_items = [
        ("Dual-Platform System Architecture: ", "Connects a Windows PC back-office workstation (Catering_Present) with a portable, 100% offline tablet/mobile kiosk (Tablet_PWA/APK) with 2-way master Excel and SQLite synchronization."),
        ("Executive Dashboard & Business KPIs: ", "Live monitoring of Gross Revenue, Total Operating Expenses, Net Profit Margins, and booking pipeline stages."),
        ("Capacity Governance & Protection: ", "Automated capacity validation preventing overbooking beyond daily kitchen production limits (e.g. max 600 pax/day)."),
        ("Customer CRM & Loyalty Engine: ", "Client booking history tracking with automated loyalty tiers (Bronze, Silver, Gold, VIP) and scheduled follow-up reminder logs.")
    ]
    for b_prefix, txt in p1_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b_prefix).bold = True
        p.add_run(txt)

    # Presenter 2
    doc.add_heading("PRESENTER 2: Client Ordering Subsystem, Sales Workflows & Billing Automation", level=2)
    p2_items = [
        ("Guided 6-Step Client Ordering Engine: ", "Client lookup & split Cebu Barangay/Street address (Step 1); touch occasion picker & Package Details Viewer (Step 2); categorized dish selection with sticky filter header (Step 3); one-tap upsell add-ons like Lechon platters (Step 4); downpayment recording & legal terms (Step 5); pre-booking confirmation drawer & automated offline PDF receipt (Step 6)."),
        ("Billing, Invoicing & Payment Records: ", "Official invoice generation (INV-XXXX), payment milestone tracking, and partial balance ledgers."),
        ("Multi-Channel Notification Gateway: ", "Automated booking confirmation and PDF invoice dispatch via Email (SMTP) and SMS (Semaphore API)."),
        ("Downpayment Enforcement: ", "Configurable threshold (e.g. 30%) blocking event confirmation without verified deposit.")
    ]
    for b_prefix, txt in p2_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b_prefix).bold = True
        p.add_run(txt)

    # Presenter 3
    doc.add_heading("PRESENTER 3: Financial Management, Expense Tracking, AI Forecasting & Kitchen Operations", level=2)
    p3_items = [
        ("Cash Flow & Expense Tracking: ", "Real-time Cash In / Cash Out monitoring, categorized expense logging (Food Costs, Labor Wages, Logistics, Utilities, Equipment), and automated Net Profit deduction (Revenue - Expenses = Net Profit)."),
        ("AI Intelligence & Demand Forecasting: ", "AI forecasting of seasonal booking demand, smart raw ingredient portion estimator per pax count, and menu profitability analytics."),
        ("Kitchen Operations (Kanban Workflow): ", "Real-time kitchen order pipeline (Pending, Cooking, Plating, Ready for Transport, Dispatched) and interactive per-dish task checklists ensuring recipe quality."),
        ("Reports & Analytics: ", "Monthly/Annual revenue comparisons, top packages report, popular dishes breakdown, and payment method distribution charts."),
        ("Audit Governance & Security: ", "Complete change audit log recording actor, action, timestamp, and JSON data states with 100% offline data privacy.")
    ]
    for b_prefix, txt in p3_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b_prefix).bold = True
        p.add_run(txt)

    # 4. Module Matrix Table
    doc.add_heading("4. Complete System Functional Module Matrix", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Module"
    hdr_cells[1].text = "Core Functionality"
    hdr_cells[2].text = "Platform"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(r'<w:shd {} w:fill="E11D48"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)

    modules = [
        ("Executive Dashboard", "Real-time KPIs, Revenue, Net Profit, Capacity alerts, Recent Activity", "PC Back-Office"),
        ("Booking & Kiosk Wizard", "6-step guided ordering, package selection, downpayment calculation", "PC + Tablet Kiosk"),
        ("Calendar & Scheduling", "Visual event timeline, date conflict detection, daily capacity load", "PC Back-Office"),
        ("Menu & Package Catalog", "CRUD packages, pax price matrix, dish categories, photo management", "PC + Tablet Kiosk"),
        ("Billing & Invoicing", "Invoice generation, downpayment enforcement, balance collection", "PC + Tablet Kiosk"),
        ("Receipt & Notifications", "PDF receipt generation, SMTP email dispatcher, SMS gateway", "PC + Tablet (PDF)"),
        ("Cash Flow Tracker", "Inflow/Outflow tracking, payment method breakdown, net balance", "PC Back-Office"),
        ("Expense & COGS Manager", "Food costs, labor wages, logistics, utilities, Net Profit margin", "PC Back-Office"),
        ("AI Demand & Forecaster", "Pax forecasting, ingredient quantity estimator, menu insights", "PC Back-Office"),
        ("Kitchen Kanban & Checklist", "Visual preparation cards, per-dish task checklist, order dispatch", "PC Back-Office"),
        ("Customer CRM & Loyalty", "Booking history, automated loyalty tiers (Bronze to VIP), follow-ups", "PC + Tablet Kiosk"),
        ("Reports & Analytics", "Revenue charts, package sales breakdown, PDF/Excel export", "PC Back-Office"),
        ("Security, Audit & Backups", "Full change audit logs, 1-click DB backup/restore, Excel data sync", "PC + Tablet Kiosk")
    ]
    for m, desc, plat in modules:
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = desc
        row_cells[2].text = plat

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. Technical Specs
    doc.add_heading("5. Technical Specifications & Architecture", level=1)
    specs = [
        "• Back-Office PC System: Native Python/Qt Desktop Application compiled to Windows executable (.exe).",
        "• Tablet Kiosk App: HTML5 / JavaScript / CSS Design System + Client-Side SQLite WASM + Standalone Android APK (7.1 MB).",
        "• 100% Offline Capability: Self-contained database engine; zero internet connection required for ordering, querying, or receipt printing.",
        "• Image Engine: Full HD 1920x1440 bicubic image processor with WebP/JPEG compression.",
        "• Data Security & Ownership: Stored locally on physical business hardware; zero third-party cloud data leaks or monthly subscription fees."
    ]
    for s in specs:
        doc.add_paragraph(s)

    doc.save(DOCX_OUT)
    print(f"Generated DOCX: {DOCX_OUT}")

    # 3. Generate PDF
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            # Header
            self.drawString(42, 755, "Jayraldine's Catering Management System — Formal Project Proposal")
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(42, 748, 570, 748)
            # Footer
            self.drawString(42, 32, "Confidential — Formal Project Proposal")
            self.drawRightString(570, 32, f"Page {self._pageNumber} of {page_count}")
            self.line(42, 42, 570, 42)
            self.restoreState()

    pdf = SimpleDocTemplate(
        PDF_OUT,
        pagesize=letter,
        leftMargin=42,
        rightMargin=42,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=17, leading=21, textColor=colors.HexColor('#E11D48'), alignment=1, spaceAfter=3)
    subtitle_style = ParagraphStyle('DocSubTitle', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=13, textColor=colors.HexColor('#64748B'), alignment=1, spaceAfter=12)
    h1_style = ParagraphStyle('H1_Custom', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor('#0F172A'), spaceBefore=10, spaceAfter=5)
    h2_style = ParagraphStyle('H2_Custom', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=10, leading=13, textColor=colors.HexColor('#E11D48'), spaceBefore=6, spaceAfter=3)
    body_style = ParagraphStyle('Body_Custom', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11.5, textColor=colors.HexColor('#1E293B'), spaceAfter=3)
    bullet_style = ParagraphStyle('Bullet_Custom', parent=body_style, leftIndent=12, bulletIndent=3, spaceAfter=2.5)

    story = []

    story.append(Paragraph("JAYRALDINE'S CATERING MANAGEMENT SYSTEM", title_style))
    story.append(Paragraph("Formal Project Proposal & Functional Specification Document", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#E11D48'), spaceAfter=8))

    # 1. Background
    story.append(Paragraph("1. Project Background & Business Case", h1_style))
    story.append(Paragraph("<b>Client:</b> Jayraldine's Catering Services (Cebu, Philippines) | <b>Classification:</b> Integrated Catering ERP & Kiosk System", body_style))
    story.append(Paragraph("<b>System Scope:</b> Dual-Platform Solution (Windows PC Back-Office Suite + 100% Offline Tablet/Mobile Field Kiosk)", body_style))
    story.append(Paragraph("<b>Project Goal:</b> Automate and modernize end-to-end catering operations — replacing manual paperwork with automated client ordering, live financial estimations, downpayment enforcement, kitchen preparation workflows, expense tracking, AI business forecasting, and comprehensive reporting.", body_style))

    story.append(Spacer(1, 6))

    # 2. Problems vs Solutions
    story.append(Paragraph("2. Operational Problems vs. Proposed System Solutions", h1_style))
    prob_pdf_data = [
        [
            Paragraph("<b>Current Manual Operational Challenge</b>", ParagraphStyle('TH', parent=body_style, textColor=colors.white, fontName='Helvetica-Bold')),
            Paragraph("<b>Proposed System Solution</b>", ParagraphStyle('TH', parent=body_style, textColor=colors.white, fontName='Helvetica-Bold'))
        ]
    ]
    for ch, sol in prob_data:
        prob_pdf_data.append([
            Paragraph(f"<b>{ch}</b>", body_style),
            Paragraph(sol, body_style)
        ])
    t_prob = Table(prob_pdf_data, colWidths=[200, 330])
    t_prob.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E11D48')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(t_prob)

    story.append(Spacer(1, 6))

    # 3. Presenter Breakdown
    story.append(Paragraph("3. Proposal Topic Breakdown for 3 Presenters", h1_style))

    # Presenter 1
    story.append(Paragraph("PRESENTER 1: Project Scope, System Architecture & Executive Oversight", h2_style))
    story.append(Paragraph("• <b>Dual-Platform System Architecture:</b> Connects a Windows PC back-office workstation (Catering_Present) with a portable, 100% offline tablet/mobile kiosk (Tablet_PWA/APK) with 2-way master Excel and SQLite synchronization.", bullet_style))
    story.append(Paragraph("• <b>Executive Dashboard & Business KPIs:</b> Live monitoring of Gross Revenue, Total Operating Expenses, Net Profit Margins, and booking pipeline stages.", bullet_style))
    story.append(Paragraph("• <b>Capacity Governance & Protection:</b> Automated capacity validation preventing overbooking beyond daily kitchen production limits (e.g. max 600 pax/day).", bullet_style))
    story.append(Paragraph("• <b>Customer CRM & Loyalty Engine:</b> Client booking history tracking with automated loyalty tiers (Bronze, Silver, Gold, VIP) and scheduled follow-up reminder logs.", bullet_style))

    story.append(Spacer(1, 4))

    # Presenter 2
    story.append(Paragraph("PRESENTER 2: Client Ordering Subsystem, Sales Workflows & Billing Automation", h2_style))
    story.append(Paragraph("• <b>Guided 6-Step Client Ordering Engine:</b> Client lookup & split Cebu Barangay/Street address (Step 1); touch occasion picker & Package Details Viewer (Step 2); categorized dish selection with sticky filter header (Step 3); one-tap upsell add-ons like Lechon platters (Step 4); downpayment recording & legal terms (Step 5); pre-booking confirmation drawer & automated offline PDF receipt (Step 6).", bullet_style))
    story.append(Paragraph("• <b>Billing, Invoicing & Payment Records:</b> Official commercial invoice generation (INV-XXXX), payment milestone tracking, and partial balance ledgers.", bullet_style))
    story.append(Paragraph("• <b>Multi-Channel Notification Gateway:</b> Automated booking confirmation and PDF invoice dispatch via Email (SMTP) and SMS (Semaphore API).", bullet_style))
    story.append(Paragraph("• <b>Downpayment Enforcement:</b> Configurable threshold (e.g. 30%) blocking event confirmation without verified deposit.", bullet_style))

    story.append(Spacer(1, 4))

    # Presenter 3
    story.append(Paragraph("PRESENTER 3: Financial Management, Expense Tracking, AI Forecasting & Kitchen Operations", h2_style))
    story.append(Paragraph("• <b>Cash Flow & Expense Tracking:</b> Real-time Cash In / Cash Out monitoring, categorized expense logging (Food Costs, Labor Wages, Logistics, Utilities, Equipment), and automated Net Profit deduction (Revenue - Expenses = Net Profit).", bullet_style))
    story.append(Paragraph("• <b>AI Intelligence & Demand Forecasting:</b> AI forecasting of seasonal booking demand, smart raw ingredient portion estimator per pax count, and menu profitability analytics.", bullet_style))
    story.append(Paragraph("• <b>Kitchen Operations (Kanban Workflow):</b> Real-time kitchen order pipeline (Pending, Cooking, Plating, Ready for Transport, Dispatched) and interactive per-dish task checklists ensuring recipe quality.", bullet_style))
    story.append(Paragraph("• <b>Reports & Analytics:</b> Monthly/Annual revenue comparisons, top packages report, popular dishes breakdown, and payment method distribution charts.", bullet_style))
    story.append(Paragraph("• <b>Audit Governance & Security:</b> Complete change audit log recording actor, action, timestamp, and JSON data states with 100% offline data privacy.", bullet_style))

    story.append(Spacer(1, 6))

    # 4. Module Matrix
    story.append(Paragraph("4. Complete System Functional Module Matrix", h1_style))
    pdf_table_data = [
        [
            Paragraph("<b>Module</b>", ParagraphStyle('TH', parent=body_style, textColor=colors.white, fontName='Helvetica-Bold')),
            Paragraph("<b>Core Functionality</b>", ParagraphStyle('TH', parent=body_style, textColor=colors.white, fontName='Helvetica-Bold')),
            Paragraph("<b>Platform</b>", ParagraphStyle('TH', parent=body_style, textColor=colors.white, fontName='Helvetica-Bold'))
        ]
    ]
    for m, desc, plat in modules:
        pdf_table_data.append([
            Paragraph(f"<b>{m}</b>", body_style),
            Paragraph(desc, body_style),
            Paragraph(plat, body_style)
        ])

    t_mod = Table(pdf_table_data, colWidths=[110, 320, 100])
    t_mod.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E11D48')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(t_mod)

    story.append(Spacer(1, 6))

    # 5. Tech Architecture
    story.append(Paragraph("5. Technical Specifications & Architecture", h1_style))
    story.append(Paragraph("• <b>Back-Office PC System:</b> Native Python/Qt Desktop Application compiled to Windows executable (.exe).", bullet_style))
    story.append(Paragraph("• <b>Tablet Kiosk App:</b> HTML5 / JavaScript / CSS + Client-Side SQLite WASM + Standalone Android APK (7.1 MB).", bullet_style))
    story.append(Paragraph("• <b>100% Offline Capability:</b> Self-contained database engine; zero internet connection required for ordering, querying, or receipt printing.", bullet_style))
    story.append(Paragraph("• <b>Data Security & Ownership:</b> Stored locally on physical business hardware; zero third-party cloud data leaks or monthly subscription fees.", bullet_style))

    pdf.build(story, canvasmaker=NumberedCanvas)
    print(f"Generated PDF: {PDF_OUT}")

if __name__ == "__main__":
    generate_all()
